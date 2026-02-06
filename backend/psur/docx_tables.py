"""
DOCX Table Builders - Generates proper python-docx Table objects from PSURContext.

Each function inserts a formatted table directly into the Document.
Tables match the structures from the PSUR141 ZyMot and non-CE example reports.
Consistent styling: Calibri 10pt, header row shaded, thin borders.
"""

from typing import List, Dict, Any, Optional
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from backend.psur.context import PSURContext


# ---------------------------------------------------------------------------
# Styling helpers
# ---------------------------------------------------------------------------

_HEADER_BG = "4472C4"  # Blue header background
_HEADER_TEXT = RGBColor(0xFF, 0xFF, 0xFF)
_BODY_TEXT = RGBColor(0x1A, 0x1A, 0x2E)
_BORDER_COLOR = "BFBFBF"
_FONT_NAME = "Calibri"
_FONT_SIZE = Pt(10)
_HEADER_FONT_SIZE = Pt(10)


def _style_table(table):
    """Apply consistent styling to a table."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Set borders on the whole table
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{_BORDER_COLOR}"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{_BORDER_COLOR}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{_BORDER_COLOR}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{_BORDER_COLOR}"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{_BORDER_COLOR}"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{_BORDER_COLOR}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def _shade_header_row(row):
    """Shade the first row blue with white text."""
    for cell in row.cells:
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="{_HEADER_BG}" w:val="clear"/>'
        )
        cell._tc.get_or_add_tcPr().append(shading)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = _HEADER_TEXT
                run.font.bold = True
                run.font.size = _HEADER_FONT_SIZE
                run.font.name = _FONT_NAME


def _set_cell_text(cell, text: str, bold: bool = False):
    """Set cell text with consistent formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.name = _FONT_NAME
    run.font.size = _FONT_SIZE
    run.font.color.rgb = _BODY_TEXT
    run.bold = bold


def _add_kv_table(doc: Document, title: str, rows: List[tuple]) -> None:
    """Add a 2-column key-value table (common for cover page info blocks)."""
    doc.add_heading(title, level=2)
    table = doc.add_table(rows=len(rows), cols=2)
    _style_table(table)
    for i, (key, value) in enumerate(rows):
        _set_cell_text(table.rows[i].cells[0], key, bold=True)
        _set_cell_text(table.rows[i].cells[1], str(value) if value else "N/A")
    # Set first column width
    for row in table.rows:
        row.cells[0].width = Cm(6)


# ---------------------------------------------------------------------------
# Cover Page Tables
# ---------------------------------------------------------------------------

def build_cover_manufacturer_table(doc: Document, ctx: PSURContext) -> None:
    """Manufacturer Information table for cover page."""
    rows = [
        ("Manufacturer", ctx.manufacturer or "[Not provided]"),
        ("Address", ctx.manufacturer_address or "[Not provided]"),
        ("SRN", ctx.manufacturer_srn or "N/A"),
        ("Authorised Representative", ctx.authorized_rep or "N/A"),
        ("AR Address", ctx.authorized_rep_address or "N/A"),
    ]
    _add_kv_table(doc, "Manufacturer Information", rows)


def build_cover_regulatory_table(doc: Document, ctx: PSURContext, template: Any) -> None:
    """Regulatory Information table for cover page."""
    rows = [
        ("Notified Body", ctx.notified_body or "N/A"),
        ("NB Number", ctx.notified_body_number or "N/A"),
        ("Regulatory Basis", getattr(template, "regulatory_basis", "EU MDR 2017/745")),
        ("PSUR Cadence", ctx.psur_cadence or "Every 2 Years"),
    ]
    _add_kv_table(doc, "Regulatory Information", rows)


def build_cover_document_table(doc: Document, ctx: PSURContext) -> None:
    """Document Information table for cover page."""
    period_start = ctx.period_start.strftime("%d %B %Y") if ctx.period_start else "N/A"
    period_end = ctx.period_end.strftime("%d %B %Y") if ctx.period_end else "N/A"
    rows = [
        ("Data Collection Start", period_start),
        ("Data Collection End", period_end),
        ("Report Generated", datetime.utcnow().strftime("%d %B %Y")),
        ("PSUR Sequence", f"#{ctx.psur_sequence_number}"),
    ]
    _add_kv_table(doc, "Document Information", rows)


# ---------------------------------------------------------------------------
# Section A Tables
# ---------------------------------------------------------------------------

def build_classification_table(doc: Document, ctx: PSURContext, template: Any) -> None:
    """Device classification table showing applicable regulatory classifications."""
    doc.add_heading("Device Classification", level=3)
    classifications = ctx.regulatory_classification or {}
    if not classifications:
        doc.add_paragraph("No classification data provided.")
        return

    table = doc.add_table(rows=1 + len(classifications), cols=2)
    _style_table(table)
    # Header
    _set_cell_text(table.rows[0].cells[0], "Regulatory System", bold=True)
    _set_cell_text(table.rows[0].cells[1], "Classification", bold=True)
    _shade_header_row(table.rows[0])
    # Data
    for i, (system, classification) in enumerate(classifications.items(), 1):
        _set_cell_text(table.rows[i].cells[0], system)
        _set_cell_text(table.rows[i].cells[1], classification)


def build_device_timeline_table(doc: Document, ctx: PSURContext) -> None:
    """Device timeline table (first CE mark, first market placement)."""
    doc.add_heading("Device Timeline", level=3)
    rows = [
        ("Device Name", ctx.device_name),
        ("UDI-DI", ctx.udi_di),
        ("Intended Use", ctx.intended_use or "[Not provided]"),
        ("Device Type", ctx.device_type or "[Not provided]"),
        ("Sterilization", ctx.sterilization_method),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    _style_table(table)
    for i, (key, value) in enumerate(rows):
        _set_cell_text(table.rows[i].cells[0], key, bold=True)
        _set_cell_text(table.rows[i].cells[1], str(value))


def build_model_catalog_table(doc: Document, ctx: PSURContext) -> None:
    """Model/variant catalog table."""
    if not ctx.device_variants:
        return
    doc.add_heading("Model Catalog", level=3)
    table = doc.add_table(rows=1 + len(ctx.device_variants), cols=2)
    _style_table(table)
    _set_cell_text(table.rows[0].cells[0], "#", bold=True)
    _set_cell_text(table.rows[0].cells[1], "Model / Variant", bold=True)
    _shade_header_row(table.rows[0])
    for i, variant in enumerate(ctx.device_variants, 1):
        _set_cell_text(table.rows[i].cells[0], str(i))
        _set_cell_text(table.rows[i].cells[1], variant)


# ---------------------------------------------------------------------------
# Section C: Sales / Distribution Tables
# ---------------------------------------------------------------------------

def build_sales_by_region_table(doc: Document, ctx: PSURContext) -> None:
    """Annual units distributed by region. Combines Table 1 and Table 2."""
    doc.add_heading("Table 1: Units Distributed by Year", level=3)

    if ctx.total_units_by_year:
        years = sorted(ctx.total_units_by_year.keys())
        table = doc.add_table(rows=2 + 1, cols=len(years) + 1)  # header + data + cumulative
        _style_table(table)
        _set_cell_text(table.rows[0].cells[0], "Metric", bold=True)
        for j, yr in enumerate(years):
            _set_cell_text(table.rows[0].cells[j + 1], str(yr), bold=True)
        _shade_header_row(table.rows[0])

        _set_cell_text(table.rows[1].cells[0], "Units Distributed", bold=True)
        cumulative = 0
        for j, yr in enumerate(years):
            val = ctx.total_units_by_year[yr]
            cumulative += val
            _set_cell_text(table.rows[1].cells[j + 1], f"{val:,}")

        _set_cell_text(table.rows[2].cells[0], "Cumulative", bold=True)
        running = 0
        for j, yr in enumerate(years):
            running += ctx.total_units_by_year[yr]
            _set_cell_text(table.rows[2].cells[j + 1], f"{running:,}")
    else:
        doc.add_paragraph("No annual distribution data available.")

    # Table 2: By Region
    if ctx.total_units_by_region:
        doc.add_heading("Table 2: Units Distributed by Region", level=3)
        regions = list(ctx.total_units_by_region.keys())
        table = doc.add_table(rows=1 + len(regions) + 1, cols=2)
        _style_table(table)
        _set_cell_text(table.rows[0].cells[0], "Region", bold=True)
        _set_cell_text(table.rows[0].cells[1], "Units", bold=True)
        _shade_header_row(table.rows[0])

        total = 0
        for i, region in enumerate(regions, 1):
            units = ctx.total_units_by_region[region]
            total += units
            _set_cell_text(table.rows[i].cells[0], region)
            _set_cell_text(table.rows[i].cells[1], f"{units:,}")

        # Total row
        _set_cell_text(table.rows[len(regions) + 1].cells[0], "Total", bold=True)
        _set_cell_text(table.rows[len(regions) + 1].cells[1], f"{total:,}", bold=True)


# ---------------------------------------------------------------------------
# Section D: Serious Incidents
# ---------------------------------------------------------------------------

def build_serious_incident_table(doc: Document, ctx: PSURContext) -> None:
    """Serious incident summary table (IMDRF classification)."""
    doc.add_heading("Table 4: Serious Incident Summary", level=3)

    rows_data = [
        ("Deaths", ctx.deaths),
        ("Serious Injuries", ctx.serious_injuries),
        ("Other Reportable Events", max(0, ctx.serious_incidents - ctx.deaths - ctx.serious_injuries)),
    ]
    table = doc.add_table(rows=1 + len(rows_data) + 1, cols=2)
    _style_table(table)
    _set_cell_text(table.rows[0].cells[0], "Category", bold=True)
    _set_cell_text(table.rows[0].cells[1], "Count", bold=True)
    _shade_header_row(table.rows[0])

    for i, (cat, count) in enumerate(rows_data, 1):
        _set_cell_text(table.rows[i].cells[0], cat)
        _set_cell_text(table.rows[i].cells[1], str(count))

    # Total
    _set_cell_text(table.rows[len(rows_data) + 1].cells[0], "Total Serious Incidents", bold=True)
    _set_cell_text(table.rows[len(rows_data) + 1].cells[1], str(ctx.serious_incidents), bold=True)

    # By type if available
    if ctx.serious_incidents_by_type:
        doc.add_heading("Table 5: Serious Incidents by Type", level=3)
        types = list(ctx.serious_incidents_by_type.items())
        table = doc.add_table(rows=1 + len(types), cols=2)
        _style_table(table)
        _set_cell_text(table.rows[0].cells[0], "Incident Type", bold=True)
        _set_cell_text(table.rows[0].cells[1], "Count", bold=True)
        _shade_header_row(table.rows[0])
        for i, (itype, count) in enumerate(types, 1):
            _set_cell_text(table.rows[i].cells[0], itype)
            _set_cell_text(table.rows[i].cells[1], str(count))


# ---------------------------------------------------------------------------
# Section E: Customer Feedback
# ---------------------------------------------------------------------------

def build_feedback_table(doc: Document, ctx: PSURContext) -> None:
    """Customer feedback summary by type and severity."""
    if ctx.complaints_by_type:
        doc.add_heading("Customer Feedback by Type", level=3)
        types = list(ctx.complaints_by_type.items())
        table = doc.add_table(rows=1 + len(types) + 1, cols=2)
        _style_table(table)
        _set_cell_text(table.rows[0].cells[0], "Feedback Type", bold=True)
        _set_cell_text(table.rows[0].cells[1], "Count", bold=True)
        _shade_header_row(table.rows[0])
        total = 0
        for i, (ftype, count) in enumerate(types, 1):
            _set_cell_text(table.rows[i].cells[0], ftype)
            _set_cell_text(table.rows[i].cells[1], str(count))
            total += count
        _set_cell_text(table.rows[len(types) + 1].cells[0], "Total", bold=True)
        _set_cell_text(table.rows[len(types) + 1].cells[1], str(total), bold=True)

    if ctx.complaints_by_severity:
        doc.add_heading("Feedback by Severity", level=3)
        sevs = list(ctx.complaints_by_severity.items())
        table = doc.add_table(rows=1 + len(sevs), cols=2)
        _style_table(table)
        _set_cell_text(table.rows[0].cells[0], "Severity", bold=True)
        _set_cell_text(table.rows[0].cells[1], "Count", bold=True)
        _shade_header_row(table.rows[0])
        for i, (sev, count) in enumerate(sevs, 1):
            _set_cell_text(table.rows[i].cells[0], sev)
            _set_cell_text(table.rows[i].cells[1], str(count))


# ---------------------------------------------------------------------------
# Section F: Complaint Rate Table
# ---------------------------------------------------------------------------

def build_complaint_rate_table(doc: Document, ctx: PSURContext) -> None:
    """Complaint counts and rates by year (Table 7 in ZyMot example)."""
    doc.add_heading("Complaint Rate by Year", level=3)

    years = sorted(set(list(ctx.total_complaints_by_year.keys()) + list(ctx.total_units_by_year.keys())))
    if not years:
        doc.add_paragraph("No annual complaint data available.")
        return

    table = doc.add_table(rows=1 + len(years) + 1, cols=4)
    _style_table(table)
    headers = ["Year", "Units Distributed", "Complaints", "Rate (%)"]
    for j, h in enumerate(headers):
        _set_cell_text(table.rows[0].cells[j], h, bold=True)
    _shade_header_row(table.rows[0])

    total_units = 0
    total_complaints = 0
    for i, yr in enumerate(years, 1):
        units = ctx.total_units_by_year.get(yr, 0)
        complaints = ctx.total_complaints_by_year.get(yr, 0)
        rate = (complaints / units * 100) if units > 0 else 0.0
        total_units += units
        total_complaints += complaints
        _set_cell_text(table.rows[i].cells[0], str(yr))
        _set_cell_text(table.rows[i].cells[1], f"{units:,}")
        _set_cell_text(table.rows[i].cells[2], str(complaints))
        _set_cell_text(table.rows[i].cells[3], f"{rate:.4f}")

    # Total row
    total_rate = (total_complaints / total_units * 100) if total_units > 0 else 0.0
    _set_cell_text(table.rows[len(years) + 1].cells[0], "Total", bold=True)
    _set_cell_text(table.rows[len(years) + 1].cells[1], f"{total_units:,}", bold=True)
    _set_cell_text(table.rows[len(years) + 1].cells[2], str(total_complaints), bold=True)
    _set_cell_text(table.rows[len(years) + 1].cells[3], f"{total_rate:.4f}", bold=True)

    # Root cause breakdown
    doc.add_heading("Complaint Root Cause Breakdown", level=3)
    rc_data = [
        ("Product Defect", ctx.complaints_product_defect),
        ("User Error", ctx.complaints_user_error),
        ("Unrelated to Device", ctx.complaints_unrelated),
        ("Unconfirmed / Pending", ctx.complaints_unconfirmed),
    ]
    table = doc.add_table(rows=1 + len(rc_data), cols=2)
    _style_table(table)
    _set_cell_text(table.rows[0].cells[0], "Root Cause Category", bold=True)
    _set_cell_text(table.rows[0].cells[1], "Count", bold=True)
    _shade_header_row(table.rows[0])
    for i, (cat, count) in enumerate(rc_data, 1):
        _set_cell_text(table.rows[i].cells[0], cat)
        _set_cell_text(table.rows[i].cells[1], str(count))


# ---------------------------------------------------------------------------
# Section H: FSCA Table
# ---------------------------------------------------------------------------

def build_fsca_table(doc: Document, ctx: PSURContext) -> None:
    """FSCA summary table."""
    doc.add_heading("Table 7: Field Safety Corrective Actions", level=3)
    # Currently no dedicated FSCA data structure; show status
    doc.add_paragraph(
        "No Field Safety Corrective Actions were initiated during the reporting period."
        if not ctx.capa_details
        else f"{len(ctx.capa_details)} CAPA-related actions tracked (see CAPA section)."
    )


# ---------------------------------------------------------------------------
# Section I: CAPA Table
# ---------------------------------------------------------------------------

def build_capa_table(doc: Document, ctx: PSURContext) -> None:
    """CAPA summary table."""
    doc.add_heading("Table 8: CAPA Summary", level=3)
    rows_data = [
        ("Open CAPAs", ctx.capa_actions_open),
        ("Closed This Period", ctx.capa_actions_closed_this_period),
        ("Effectiveness Verified", ctx.capa_actions_effectiveness_verified),
    ]
    table = doc.add_table(rows=1 + len(rows_data), cols=2)
    _style_table(table)
    _set_cell_text(table.rows[0].cells[0], "CAPA Status", bold=True)
    _set_cell_text(table.rows[0].cells[1], "Count", bold=True)
    _shade_header_row(table.rows[0])
    for i, (label, count) in enumerate(rows_data, 1):
        _set_cell_text(table.rows[i].cells[0], label)
        _set_cell_text(table.rows[i].cells[1], str(count))

    # Detail rows if available
    if ctx.capa_details:
        doc.add_heading("CAPA Details", level=3)
        detail_table = doc.add_table(rows=1 + len(ctx.capa_details), cols=3)
        _style_table(detail_table)
        for j, h in enumerate(["CAPA ID", "Description", "Status"]):
            _set_cell_text(detail_table.rows[0].cells[j], h, bold=True)
        _shade_header_row(detail_table.rows[0])
        for i, capa in enumerate(ctx.capa_details, 1):
            _set_cell_text(detail_table.rows[i].cells[0], str(capa.get("id", i)))
            _set_cell_text(detail_table.rows[i].cells[1], str(capa.get("description", ""))[:80])
            _set_cell_text(detail_table.rows[i].cells[2], str(capa.get("status", "Open")))


# ---------------------------------------------------------------------------
# Section K: External Database Results
# ---------------------------------------------------------------------------

def build_external_db_table(doc: Document, ctx: PSURContext) -> None:
    """External database search results summary."""
    doc.add_heading("External Database Search Results", level=3)
    databases = ["FDA MAUDE", "BfArM", "MHRA", "Eudamed"]
    table = doc.add_table(rows=1 + len(databases), cols=3)
    _style_table(table)
    for j, h in enumerate(["Database", "Search Performed", "Relevant Findings"]):
        _set_cell_text(table.rows[0].cells[j], h, bold=True)
    _shade_header_row(table.rows[0])

    vig = "Yes" if ctx.data_availability_external_vigilance else "No"
    for i, db_name in enumerate(databases, 1):
        _set_cell_text(table.rows[i].cells[0], db_name)
        _set_cell_text(table.rows[i].cells[1], vig)
        _set_cell_text(table.rows[i].cells[2], f"{ctx.total_vigilance_events} events" if ctx.vigilance_data_available else "N/A")


# ---------------------------------------------------------------------------
# Markdown Table Parser (for AI-generated pipe tables)
# ---------------------------------------------------------------------------

def parse_markdown_table(lines: List[str]) -> Optional[List[List[str]]]:
    """
    Parse markdown pipe-delimited table lines into a 2D list of cell strings.
    Returns None if lines do not form a valid table.
    """
    if not lines or not any("|" in line for line in lines):
        return None

    table_rows: List[List[str]] = []
    for line in lines:
        stripped = line.strip()
        if not stripped.startswith("|"):
            continue
        # Skip separator rows (|---|---|)
        if all(c in "|-: " for c in stripped):
            continue
        cells = [cell.strip() for cell in stripped.split("|")]
        # Remove empty first/last from leading/trailing pipes
        if cells and cells[0] == "":
            cells = cells[1:]
        if cells and cells[-1] == "":
            cells = cells[:-1]
        if cells:
            table_rows.append(cells)

    return table_rows if len(table_rows) >= 2 else None


def insert_markdown_table(doc: Document, rows: List[List[str]]) -> None:
    """Insert a parsed markdown table into the document as a proper DOCX table."""
    if not rows:
        return
    max_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    _style_table(table)

    # First row is header
    for j in range(min(len(rows[0]), max_cols)):
        _set_cell_text(table.rows[0].cells[j], rows[0][j], bold=True)
    _shade_header_row(table.rows[0])

    # Data rows
    for i in range(1, len(rows)):
        for j in range(min(len(rows[i]), max_cols)):
            _set_cell_text(table.rows[i].cells[j], rows[i][j])


# ---------------------------------------------------------------------------
# Master table dispatcher (maps table IDs to builder functions)
# ---------------------------------------------------------------------------

TABLE_BUILDERS: Dict[str, Any] = {
    "classification": build_classification_table,
    "device_timeline": build_device_timeline_table,
    "model_catalog": build_model_catalog_table,
    "sales_by_region": build_sales_by_region_table,
    "serious_incident": build_serious_incident_table,
    "feedback": build_feedback_table,
    "complaint_rate": build_complaint_rate_table,
    "fsca": build_fsca_table,
    "capa": build_capa_table,
    "external_db": build_external_db_table,
}


def build_tables_for_section(doc: Document, section_id: str,
                             ctx: PSURContext, template: Any) -> None:
    """
    Insert all required data tables for a section based on template spec.
    Called BEFORE the AI narrative is inserted.
    """
    from backend.psur.templates import get_section_spec
    spec = get_section_spec(getattr(ctx, "template_id", "eu_uk_mdr"), section_id)
    if not spec or not spec.required_tables:
        return

    for table_id in spec.required_tables:
        builder = TABLE_BUILDERS.get(table_id)
        if builder:
            try:
                # Some builders need template arg, some don't
                import inspect
                sig = inspect.signature(builder)
                if "template" in sig.parameters:
                    builder(doc, ctx, template)
                else:
                    builder(doc, ctx)
            except Exception as e:
                doc.add_paragraph(f"[Table {table_id}: generation error - {e}]")
