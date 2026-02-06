"""
FastAPI Main Application
Provides REST API and WebSocket endpoints for the Multi-Agent PSUR System
"""

from fastapi import FastAPI, WebSocket, WebSocketDisconnect, HTTPException, UploadFile, File, Depends, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from sqlalchemy.orm import Session
from typing import List, Dict, Any, Optional
from pydantic import BaseModel
import asyncio
import io
import json
import traceback
from datetime import datetime

from backend.database.session import get_db_context, get_db, init_db
from backend.database.models import (
    PSURSession, Agent, ChatMessage, SectionDocument,
    WorkflowState, DataFile, ChartAsset
)
from backend.psur import SOTAOrchestrator, AGENT_ROLES, SECTION_DEFINITIONS, WORKFLOW_ORDER
from backend.psur.context import PSURContext
from backend.psur.extraction import extract_from_file
from backend.config import AGENT_CONFIGS, settings

# Initialize FastAPI
app = FastAPI(
    title="Multi-Agent PSUR System API",
    description="REST API and WebSocket for AI-powered PSUR generation",
    version="1.0.0"
)

# CORS Configuration (permissive for MVP; restrict in production)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Global exception handler to ensure CORS headers are always sent
@app.exception_handler(Exception)
async def global_exception_handler(request: Request, exc: Exception):
    traceback.print_exc()
    return JSONResponse(
        status_code=500,
        content={"detail": str(exc)},
        headers={
            "Access-Control-Allow-Origin": "*",
            "Access-Control-Allow-Methods": "*",
            "Access-Control-Allow-Headers": "*",
        }
    )

# Startup event to initialize database
@app.on_event("startup")
async def startup_event():
    print("Initializing database...")
    init_db()
    print("Database initialized successfully")

    # Check matplotlib availability at startup
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        print("[startup] matplotlib is available for chart generation")
    except ImportError:
        print("[startup] WARNING: matplotlib not installed. Charts will NOT be generated.")
        print("[startup] Install with: pip install matplotlib")

# WebSocket Connection Manager
class ConnectionManager:
    def __init__(self):
        self.active_connections: List[WebSocket] = []

    async def connect(self, websocket: WebSocket):
        await websocket.accept()
        self.active_connections.append(websocket)

    def disconnect(self, websocket: WebSocket):
        if websocket in self.active_connections:
            self.active_connections.remove(websocket)

    async def broadcast(self, message: dict):
        """Broadcast message to all connected clients"""
        for connection in self.active_connections:
            try:
                await connection.send_json(message)
            except Exception as e:
                print(f"Error broadcasting to client: {e}")

manager = ConnectionManager()

# Global orchestrator tracking for pause/resume functionality
active_orchestrators: Dict[int, SOTAOrchestrator] = {}

# ============================================================================
# REST API ENDPOINTS
# ============================================================================

@app.get("/")
async def root():
    """Health check endpoint"""
    return {
        "status": "online",
        "service": "Multi-Agent PSUR System",
        "version": "1.0.0"
    }

@app.post("/api/sessions")
async def create_session(
    device_name: str,
    udi_di: str,
    start_date: str,
    end_date: str,
    template_id: str = "eu_uk_mdr",
    db: Session = Depends(get_db)
):
    """Create a new PSUR generation session"""
    try:
        # Parse dates
        start = datetime.fromisoformat(start_date)
        end = datetime.fromisoformat(end_date)

        session = PSURSession(
            device_name=device_name,
            udi_di=udi_di,
            period_start=start,
            period_end=end,
            template_id=template_id,
            status="initializing",
            created_at=datetime.utcnow()
        )
        db.add(session)
        db.commit()
        db.refresh(session)

        # Initialize workflow state
        workflow = WorkflowState(
            session_id=session.id,
            current_section="A",
            sections_completed=0,
            total_sections=13,
            status="initialized"
        )
        db.add(workflow)
        db.commit()

        return {
            "session_id": session.id,
            "device_name": device_name,
            "template_id": template_id,
            "status": "initialized"
        }
    except Exception as e:
        db.rollback()
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/sessions")
async def list_sessions(db: Session = Depends(get_db)):
    """List all sessions"""
    try:
        sessions = db.query(PSURSession).order_by(PSURSession.created_at.desc()).all()
        
        result = []
        for session in sessions:
            workflow = db.query(WorkflowState).filter(WorkflowState.session_id == session.id).first()
            result.append({
                "id": session.id,
                "device_name": session.device_name,
                "udi_di": session.udi_di,
                "period_start": session.period_start.isoformat() if session.period_start else None,
                "period_end": session.period_end.isoformat() if session.period_end else None,
                "status": session.status,
                "created_at": session.created_at.isoformat() if session.created_at else None,
                "workflow": {
                    "sections_completed": workflow.sections_completed if workflow else 0,
                    "total_sections": workflow.total_sections if workflow else 13,
                } if workflow else None
            })
        
        return result
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

@app.delete("/api/sessions/{session_id}")
async def delete_session(session_id: int, db: Session = Depends(get_db)):
    """Delete a session and all associated data"""
    try:
        session = db.query(PSURSession).filter(PSURSession.id == session_id).first()
        if not session:
            raise HTTPException(status_code=404, detail="Session not found")
        
        # Delete related records in correct order (foreign key dependencies)
        db.query(Agent).filter(Agent.session_id == session_id).delete(synchronize_session='fetch')
        db.query(ChatMessage).filter(ChatMessage.session_id == session_id).delete(synchronize_session='fetch')
        db.query(SectionDocument).filter(SectionDocument.session_id == session_id).delete(synchronize_session='fetch')
        db.query(WorkflowState).filter(WorkflowState.session_id == session_id).delete(synchronize_session='fetch')
        db.query(DataFile).filter(DataFile.session_id == session_id).delete(synchronize_session='fetch')
        
        # Now delete the session
        db.delete(session)
        db.commit()
        
        return {"status": "success", "message": f"Session {session_id} deleted"}
    except HTTPException:
        raise
    except Exception as e:
        traceback.print_exc()
        db.rollback()
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/sessions/{session_id}")
async def get_session(session_id: int, db: Session = Depends(get_db)):
    """Get session details"""
    try:
        session = db.query(PSURSession).filter(PSURSession.id == session_id).first()
        if not session:
            raise HTTPException(status_code=404, detail="Session not found")
        
        workflow = db.query(WorkflowState).filter(WorkflowState.session_id == session_id).first()
        
        return {
            "id": session.id,
            "device_name": session.device_name,
            "udi_di": session.udi_di,
            "start_date": session.period_start.isoformat() if session.period_start else None,
            "end_date": session.period_end.isoformat() if session.period_end else None,
            "status": session.status,
            "created_at": session.created_at.isoformat() if session.created_at else None,
            "workflow": {
                "current_section": workflow.current_section if workflow else None,
                "sections_completed": workflow.sections_completed if workflow else 0,
                "total_sections": workflow.total_sections if workflow else 13,
                "status": workflow.status if workflow else "unknown"
            }
        }
    except HTTPException:
        raise
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

from backend.psur.extraction import analyze_upload

@app.post("/api/sessions/{session_id}/upload")
async def upload_file(
    session_id: int,
    file: UploadFile = File(...),
    file_type: str = "sales",
    db: Session = Depends(get_db)
):
    """Upload data file (sales, complaints, etc.)"""
    try:
        content = await file.read()
        
        # Save to database
        data_file = DataFile(
            session_id=session_id,
            file_type=file_type,
            filename=file.filename,
            file_data=content,
            uploaded_at=datetime.utcnow()
        )
        db.add(data_file)
        
        # Process file data
        result = analyze_upload(content, file.filename, file_type)
        analysis = result["summary"]
        metadata = result["metadata"]
        
        # Build column detection diagnostics message
        columns_detected = metadata.get("columns_detected", {})
        all_columns = metadata.get("all_columns", [])
        diagnostics_lines = []
        
        if file_type == "sales":
            units_col = columns_detected.get("units")
            year_col = columns_detected.get("year")
            diagnostics_lines.append(f"Units Column: {units_col if units_col else 'NOT FOUND'}")
            diagnostics_lines.append(f"Year Column: {year_col if year_col else 'NOT FOUND'}")
            if not units_col:
                diagnostics_lines.append(f"WARNING: Could not detect units/quantity column. Available: {all_columns}")
        elif file_type == "complaints":
            severity_col = columns_detected.get("severity")
            closure_col = columns_detected.get("closure")
            type_col = columns_detected.get("type")
            root_cause_col = columns_detected.get("root_cause")
            diagnostics_lines.append(f"Severity Column: {severity_col if severity_col else 'NOT FOUND'}")
            diagnostics_lines.append(f"Closure/Status Column: {closure_col if closure_col else 'NOT FOUND'}")
            diagnostics_lines.append(f"Type Column: {type_col if type_col else 'NOT FOUND'}")
            diagnostics_lines.append(f"Root Cause Column: {root_cause_col if root_cause_col else 'NOT FOUND'}")
            if not severity_col:
                diagnostics_lines.append(f"WARNING: No severity column detected. Available: {all_columns}")
        elif file_type in ("vigilance", "maude"):
            type_col = columns_detected.get("type")
            diagnostics_lines.append(f"Event Type Column: {type_col if type_col else 'NOT FOUND'}")
        
        diagnostics_section = "\n".join(diagnostics_lines) if diagnostics_lines else ""
        record_count = metadata.get("record_count", 0)
        
        # Save analysis as a system message so agents can see it
        analysis_msg = ChatMessage(
            session_id=session_id,
            from_agent="System",
            to_agent="all",
            message=f"""Data Analysis: {file_type.upper()} ({file.filename})

**Records Found:** {record_count}
**Columns Detected:**
{diagnostics_section}

{analysis}""",
            message_type="system"
        )
        db.add(analysis_msg)
        
        # If UDI-DI is found in metadata, update the session
        if "udi_di" in metadata:
            session = db.query(PSURSession).filter(PSURSession.id == session_id).first()
            if session and (not session.udi_di or session.udi_di == "Pending Extraction"):
                session.udi_di = metadata["udi_di"]
                # Also notify agents
                udi_msg = ChatMessage(
                    session_id=session_id,
                    from_agent="System",
                    to_agent="all",
                    message=f"Metadata Detected: Detected UDI-DI: {metadata['udi_di']} from {file.filename}",
                    message_type="success"
                )
                db.add(udi_msg)
        
        db.commit()
        db.refresh(data_file)
        
        return {
            "file_id": data_file.id,
            "filename": file.filename,
            "file_type": file_type,
            "size": len(content),
            "status": "uploaded",
            "analysis_preview": (analysis or "")[:100] + ("..." if len(analysis or "") > 100 else "")
        }
    except Exception as e:
        db.rollback()
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/sessions/{session_id}/files")
async def get_session_files(session_id: int, db: Session = Depends(get_db)):
    """Get all uploaded files for a session"""
    try:
        files = db.query(DataFile).filter(DataFile.session_id == session_id).all()
        
        return [
            {
                "id": f.id,
                "filename": f.filename,
                "file_type": f.file_type,
                "size": len(f.file_data) if f.file_data else 0,
                "uploaded_at": f.uploaded_at.isoformat() if f.uploaded_at else None
            }
            for f in files
        ]
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

class MasterContextIntakeBody(BaseModel):
    denominator_scope: str = "reporting_period_only"
    inference_policy: str = "strictly_factual"
    closure_definition: str = ""
    baseline_year: Optional[int] = None
    external_vigilance_searched: bool = False
    complaint_closures_complete: bool = False
    rmf_hazard_list_available: bool = False
    intended_use_provided: bool = False


@app.patch("/api/sessions/{session_id}/intake")
async def set_master_context_intake(
    session_id: int,
    body: MasterContextIntakeBody,
    db: Session = Depends(get_db),
):
    """Set master context intake options before starting PSUR generation."""
    try:
        session = db.query(PSURSession).filter(PSURSession.id == session_id).first()
        if not session:
            raise HTTPException(status_code=404, detail="Session not found")
        intake = {
            "denominator_scope": body.denominator_scope,
            "inference_policy": body.inference_policy,
            "closure_definition": body.closure_definition or None,
            "baseline_year": body.baseline_year,
            "external_vigilance_searched": body.external_vigilance_searched,
            "complaint_closures_complete": body.complaint_closures_complete,
            "rmf_hazard_list_available": body.rmf_hazard_list_available,
            "intended_use_provided": body.intended_use_provided,
        }
        setattr(session, "master_context_intake", intake)
        db.commit()
        return {"status": "ok", "session_id": session_id}
    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/sessions/{session_id}/validate")
async def validate_session_data(session_id: int, db: Session = Depends(get_db)):
    """Check if uploaded data can be parsed correctly before starting generation.
    Uses the same unified extraction.py pipeline as the orchestrator."""
    try:
        session = db.query(PSURSession).filter(PSURSession.id == session_id).first()
        if not session:
            raise HTTPException(status_code=404, detail="Session not found")

        data_files = db.query(DataFile).filter(DataFile.session_id == session_id).all()

        issues: List[Dict[str, str]] = []

        # Check for required file types
        file_types = [getattr(f, "file_type", "") for f in data_files]
        if "sales" not in file_types:
            issues.append({"severity": "warning", "message": "No sales/distribution data file uploaded. Units distributed will be 0."})
        if "complaints" not in file_types:
            issues.append({"severity": "warning", "message": "No complaints data file uploaded. Complaint analysis will be limited."})
        if len(data_files) == 0:
            issues.append({"severity": "error", "message": "No data files uploaded. Please upload at least one data file."})

        # Run the SAME extraction pipeline the orchestrator uses
        ctx = PSURContext()
        all_mappings: Dict[str, Any] = {}

        for df_obj in data_files:
            _file_type = getattr(df_obj, "file_type", "") or ""
            _filename = getattr(df_obj, "filename", "") or ""
            _file_data = getattr(df_obj, "file_data", b"") or b""

            diag = extract_from_file(_file_data, _filename, _file_type, ctx)
            if diag.get("warnings"):
                for w in diag["warnings"]:
                    issues.append({"severity": "warning", "message": w})
            if diag.get("columns_detected"):
                all_mappings[_filename] = diag["columns_detected"]

        ctx.calculate_metrics()

        # Check extraction results
        if ctx.total_units_sold == 0 and "sales" in file_types:
            issues.append({
                "severity": "error",
                "message": "Could not extract units distributed from sales files. Check column names. Looking for: units, quantity, qty, sold, distributed, shipped, volume."
            })

        if ctx.total_complaints > 0:
            if not ctx.complaints_by_severity:
                issues.append({
                    "severity": "warning",
                    "message": "Complaints found but severity column not detected. Severity analysis will be limited."
                })
            if ctx.complaints_closed_count == 0:
                issues.append({
                    "severity": "warning",
                    "message": "Complaints found but no closure/status column detected or no closed complaints. Investigation closure rate is 0%."
                })

        has_errors = any(i["severity"] == "error" for i in issues)

        return {
            "valid": not has_errors,
            "issues": issues,
            "extracted_data": {
                "total_units_sold": ctx.total_units_sold,
                "total_complaints": ctx.total_complaints,
                "complaints_closed_count": ctx.complaints_closed_count,
                "complaints_with_root_cause_identified": ctx.complaints_with_root_cause_identified,
                "total_units_by_year": ctx.total_units_by_year,
                "total_complaints_by_year": ctx.total_complaints_by_year,
                "serious_incidents": ctx.serious_incidents,
                "total_vigilance_events": ctx.total_vigilance_events,
                "has_sales": ctx.sales_data_available,
                "has_complaints": ctx.complaint_data_available,
                "has_vigilance": ctx.vigilance_data_available,
                "column_mappings": all_mappings,
            }
        }
    except HTTPException:
        raise
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Validation failed: {str(e)}")


@app.post("/api/sessions/{session_id}/start")
async def start_generation(session_id: int, db: Session = Depends(get_db)):
    """Start PSUR generation. The orchestrator handles all extraction internally."""
    try:
        session = db.query(PSURSession).filter(PSURSession.id == session_id).first()
        if not session:
            raise HTTPException(status_code=404, detail="Session not found")

        setattr(session, "status", "running")
        db.commit()

        asyncio.create_task(run_orchestrator(session_id))

        return {
            "status": "started",
            "session_id": session_id,
            "message": "PSUR generation initiated.",
        }
    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/sessions/{session_id}/messages")
async def get_messages(
    session_id: int,
    limit: int = 100,
    db: Session = Depends(get_db)
):
    """Get chat messages for a session"""
    try:
        messages = db.query(ChatMessage)\
            .filter(ChatMessage.session_id == session_id)\
            .order_by(ChatMessage.timestamp.desc())\
            .limit(limit)\
            .all()
        
        return [
            {
                "id": m.id,
                "from_agent": m.from_agent,
                "to_agent": m.to_agent,
                "message": m.message,
                "message_type": m.message_type,
                "timestamp": m.timestamp.isoformat() if m.timestamp else None
            }
            for m in reversed(messages)
        ]
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/sessions/{session_id}/agents")
async def get_agents(session_id: int, db: Session = Depends(get_db)):
    """Get all agents with their status and model info (SOTA agent definitions)"""
    try:
        # Get status from DB
        db_agents = db.query(Agent).filter(Agent.session_id == session_id).all()
        status_map = {a.agent_id: a.status for a in db_agents}

        agents_list = []
        for agent_id, agent_info in AGENT_ROLES.items():
            config = AGENT_CONFIGS.get(agent_id)
            agents_list.append({
                "id": agent_id,
                "name": agent_info["name"],
                "role": agent_info["role"],
                "title": agent_info.get("title", agent_info["role"]),
                "expertise": agent_info.get("expertise", ""),
                "primary_section": agent_info.get("primary_section"),
                "color": agent_info.get("color", "#6366f1"),
                "ai_provider": config.ai_provider if config else "anthropic",
                "model": config.model if config else "claude-sonnet-4-20250514",
                "status": status_map.get(agent_id, "idle")
            })

        return agents_list
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/sessions/{session_id}/sections")
async def get_sections(session_id: int, db: Session = Depends(get_db)):
    """Get all section documents"""
    try:
        sections = db.query(SectionDocument)\
            .filter(SectionDocument.session_id == session_id)\
            .order_by(SectionDocument.section_id)\
            .all()
        
        return [
            {
                "section_id": s.section_id,
                "section_name": s.section_name,
                "author_agent": s.author_agent,
                "status": s.status,
                "word_count": len(s.content.split()) if s.content else 0,
                "created_at": s.created_at.isoformat() if s.created_at else None,
                "qc_feedback": s.qc_feedback
            }
            for s in sections
        ]
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/sessions/{session_id}/sections/{section_id}")
async def get_section_content(
    session_id: int,
    section_id: str,
    db: Session = Depends(get_db)
):
    """Get full content of a specific section"""
    try:
        section = db.query(SectionDocument)\
            .filter(
                SectionDocument.session_id == session_id,
                SectionDocument.section_id == section_id
            )\
            .first()
        
        if not section:
            raise HTTPException(status_code=404, detail="Section not found")
        
        return {
            "section_id": section.section_id,
            "section_name": section.section_name,
            "author_agent": section.author_agent,
            "content": section.content,
            "status": section.status,
            "qc_feedback": section.qc_feedback,
            "created_at": section.created_at.isoformat() if section.created_at else None
        }
    except HTTPException:
        raise
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/sessions/{session_id}/document")
async def get_complete_document(session_id: int, db: Session = Depends(get_db)):
    """Get the complete PSUR document with all sections"""
    try:
        session = db.query(PSURSession).filter(PSURSession.id == session_id).first()
        if not session:
            raise HTTPException(status_code=404, detail="Session not found")
        
        sections = db.query(SectionDocument)\
            .filter(SectionDocument.session_id == session_id)\
            .order_by(SectionDocument.section_id)\
            .all()
        
        return {
            "session_id": session_id,
            "device_name": session.device_name,
            "udi_di": session.udi_di,
            "period_start": session.period_start.isoformat() if session.period_start else None,
            "period_end": session.period_end.isoformat() if session.period_end else None,
            "status": session.status,
            "sections": [
                {
                    "section_id": s.section_id,
                    "section_name": s.section_name,  
                    "author_agent": s.author_agent,
                    "content": s.content,
                    "word_count": len(s.content.split()) if s.content else 0
                }
                for s in sections
            ],
            "total_sections": len(sections)
        }
    except HTTPException:
        raise
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

class MessageInput(BaseModel):
    message: str
    from_agent: str = "User"
    to_agent: str = "all" 
    message_type: str = "normal"

@app.post("/api/sessions/{session_id}/messages")
async def create_message(
    session_id: int,
    input: MessageInput,
    db: Session = Depends(get_db)
):
    """Post a new message to the session. If from User, generate AI response."""
    try:
        msg = ChatMessage(
            session_id=session_id,
            from_agent=input.from_agent,
            to_agent=input.to_agent,
            message=input.message,
            message_type=input.message_type,
            timestamp=datetime.utcnow(),
            processed=False,
        )
        db.add(msg)
        db.commit()
        db.refresh(msg)

        # Notify clients via WebSocket
        await manager.broadcast({
            "type": "new_message",
            "data": {
                "id": msg.id,
                "from_agent": msg.from_agent,
                "to_agent": msg.to_agent,
                "message": msg.message,
                "message_type": msg.message_type,
                "timestamp": msg.timestamp.isoformat() if msg.timestamp else None
            }
        })

        # If from User, generate an AI response asynchronously
        if input.from_agent == "User":
            asyncio.create_task(_respond_to_user_message(
                session_id, msg.id, input.message, input.to_agent
            ))

        return {"status": "ok", "message_id": msg.id}
    except Exception as e:
        db.rollback()
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


async def _respond_to_user_message(session_id: int, msg_id: int,
                                    message: str, target: str):
    """Generate an AI response to a user message and post it to chat."""
    from backend.psur.ai_client import call_ai

    try:
        # Determine which agent should respond
        responder = "Alex"  # default
        for name in AGENT_CONFIGS:
            if f"@{name}" in message or f"@{name.lower()}" in message.lower():
                responder = name
                break
        if target != "all" and target in AGENT_CONFIGS:
            responder = target

        # Get orchestrator context if available
        ctx_summary = ""
        if session_id in active_orchestrators:
            orch = active_orchestrators[session_id]
            if orch.context:
                ctx_summary = (
                    f"Device: {orch.context.device_name}. "
                    f"Units: {orch.context.total_units_sold:,}. "
                    f"Complaints: {orch.context.total_complaints}. "
                    f"Phase: {orch.current_phase}."
                )

        cfg = AGENT_CONFIGS.get(responder, AGENT_CONFIGS.get("Alex"))
        role_info = AGENT_ROLES.get(responder, {})
        personality = role_info.get("personality", "")
        sys_prompt = (
            f"You are {responder}, {cfg.role if cfg else 'PSUR Agent'}. {personality} "
            f"{ctx_summary} "
            "Respond concisely in prose (2-4 sentences). No bullet points. "
            "You are speaking directly to the user in a professional team chat."
        )

        response = await call_ai(responder, sys_prompt, f'User says: "{message}"')
        if not response:
            response = f"I apologize, I could not generate a response at this time. Please try again."

        # Save response to DB
        with get_db_context() as db2:
            resp_msg = ChatMessage(
                session_id=session_id,
                from_agent=responder,
                to_agent="User",
                message=response,
                message_type="normal",
                timestamp=datetime.utcnow(),
                response_to_id=msg_id,
            )
            db2.add(resp_msg)

            # Mark original message as processed (Boolean column only)
            orig = db2.query(ChatMessage).filter(ChatMessage.id == msg_id).first()
            if orig:
                orig.processed = True
            db2.commit()

        # Broadcast response
        await manager.broadcast({
            "type": "new_message",
            "data": {
                "from_agent": responder,
                "to_agent": "User",
                "message": response,
                "message_type": "normal",
                "timestamp": datetime.utcnow().isoformat()
            }
        })

    except Exception as e:
        print(f"[chat] Error responding to user message: {e}")
        traceback.print_exc()


# ============================================================================
# INTERACTIVE WORKFLOW CONTROL ENDPOINTS
# ============================================================================

@app.post("/api/sessions/{session_id}/pause")
async def pause_workflow(session_id: int, db: Session = Depends(get_db)):
    """Pause the workflow at the next checkpoint."""
    try:
        if session_id not in active_orchestrators:
            raise HTTPException(status_code=404, detail="No active workflow for this session")
        
        orchestrator = active_orchestrators[session_id]
        success = orchestrator.request_pause()
        
        if success:
            # Broadcast pause status
            await manager.broadcast({
                "type": "workflow_paused",
                "session_id": session_id
            })
            return {"status": "pausing", "message": "Workflow will pause at next checkpoint"}
        else:
            return {"status": "not_running", "message": "Workflow is not currently running"}
    except HTTPException:
        raise
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/sessions/{session_id}/resume")
async def resume_workflow(session_id: int, db: Session = Depends(get_db)):
    """Resume a paused workflow."""
    try:
        if session_id not in active_orchestrators:
            raise HTTPException(status_code=404, detail="No active workflow for this session")
        
        orchestrator = active_orchestrators[session_id]
        success = orchestrator.request_resume()
        
        if success:
            # Broadcast resume status
            await manager.broadcast({
                "type": "workflow_resumed",
                "session_id": session_id
            })
            return {"status": "resumed", "message": "Workflow has resumed"}
        else:
            return {"status": "not_paused", "message": "Workflow is not currently paused"}
    except HTTPException:
        raise
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


class AskAgentInput(BaseModel):
    agent: str
    question: str


@app.post("/api/sessions/{session_id}/ask")
async def ask_agent(session_id: int, input: AskAgentInput, db: Session = Depends(get_db)):
    """Ask a specific agent a direct question."""
    try:
        # Validate agent exists
        if input.agent not in AGENT_CONFIGS:
            raise HTTPException(status_code=400, detail=f"Unknown agent: {input.agent}")
        
        # Check if session exists
        session = db.query(PSURSession).filter(PSURSession.id == session_id).first()
        if not session:
            raise HTTPException(status_code=404, detail="Session not found")
        
        # First, post the user's question to the chat
        user_msg = ChatMessage(
            session_id=session_id,
            from_agent="User",
            to_agent=input.agent,
            message=input.question,
            message_type="normal",
            timestamp=datetime.utcnow(),
            processed=False
        )
        db.add(user_msg)
        db.commit()
        db.refresh(user_msg)
        
        # Broadcast user message
        await manager.broadcast({
            "type": "new_message",
            "data": {
                "id": user_msg.id,
                "from_agent": "User",
                "to_agent": input.agent,
                "message": input.question,
                "message_type": "normal",
                "timestamp": user_msg.timestamp.isoformat() if user_msg.timestamp else None
            }
        })
        
        # Use active orchestrator if available; otherwise use a lightweight direct call
        if session_id in active_orchestrators:
            orchestrator = active_orchestrators[session_id]
            result = await orchestrator.ask_agent_directly(input.agent, input.question)
        else:
            # No active workflow -- respond directly without heavy initialization
            from backend.psur.ai_client import call_ai
            cfg = AGENT_CONFIGS.get(input.agent, AGENT_CONFIGS.get("Alex"))
            role_info = AGENT_ROLES.get(input.agent, {})
            personality = role_info.get("personality", "")
            sys_prompt = (
                f"You are {input.agent}, {cfg.role if cfg else 'PSUR Agent'}. {personality} "
                f"Device: {session.device_name}. Answer directly and concisely."
            )
            response = await call_ai(input.agent, sys_prompt, input.question)
            # Save the response to chat
            resp_msg = ChatMessage(
                session_id=session_id,
                from_agent=input.agent,
                to_agent="User",
                message=response or "No response available.",
                message_type="normal",
                timestamp=datetime.utcnow(),
                response_to_id=user_msg.id,
            )
            db.add(resp_msg)
            result = {"response": response, "agent": input.agent, "error": None}
        
        # Mark user message as processed
        user_msg.processed = True
        db.commit()
        
        if result.get("error"):
            return {
                "status": "error",
                "error": result["error"],
                "agent": input.agent
            }
        
        return {
            "status": "ok",
            "response": result.get("response"),
            "agent": input.agent
        }
    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/sessions/{session_id}/workflow-status")
async def get_workflow_status(session_id: int, db: Session = Depends(get_db)):
    """Get current workflow status including pause state."""
    try:
        session = db.query(PSURSession).filter(PSURSession.id == session_id).first()
        if not session:
            # Return safe default instead of 404 -- the frontend polls this endpoint
            return {
                "session_id": session_id,
                "session_status": "unknown",
                "orchestrator_active": False,
                "workflow_state": None,
                "orchestrator_status": None
            }
        
        workflow_state = db.query(WorkflowState).filter(
            WorkflowState.session_id == session_id
        ).first()
        
        orchestrator_active = session_id in active_orchestrators
        orchestrator_status = None
        
        if orchestrator_active:
            try:
                orchestrator_status = active_orchestrators[session_id].get_workflow_status()
            except Exception:
                pass
        
        ws_data = None
        if workflow_state:
            ws_data = {
                "status": getattr(workflow_state, "status", "not_started"),
                "current_section": getattr(workflow_state, "current_section", None),
                "sections_completed": getattr(workflow_state, "sections_completed", 0),
                "total_sections": getattr(workflow_state, "total_sections", 13),
                "paused": getattr(workflow_state, "paused", False),
                "current_agent": getattr(workflow_state, "current_agent", None),
            }
        
        return {
            "session_id": session_id,
            "session_status": session.status,
            "orchestrator_active": orchestrator_active,
            "workflow_state": ws_data,
            "orchestrator_status": orchestrator_status
        }
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/agents")
async def list_agents():
    """List all available agents and their roles."""
    agents = []
    for name, config in AGENT_CONFIGS.items():
        agents.append({
            "name": name,
            "role": config.role,
            "ai_provider": config.ai_provider,
            "model": config.model
        })
    return {"agents": agents}


@app.get("/api/templates")
async def list_templates():
    """List available PSUR report templates."""
    from backend.psur.templates import get_template_choices
    return {"templates": get_template_choices()}


@app.get("/api/sessions/{session_id}/document/download")
async def download_document(session_id: int, db: Session = Depends(get_db)):
    """Download complete PSUR as DOCX with rich cover page, data tables, and inline charts."""
    try:
        from fastapi.responses import FileResponse
        from docx import Document as DocxDocument
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import tempfile
        import os
        import re

        from backend.psur.templates import load_template
        from backend.psur.docx_tables import (
            build_cover_manufacturer_table,
            build_cover_regulatory_table,
            build_cover_document_table,
            build_tables_for_section,
            parse_markdown_table,
            insert_markdown_table,
        )
        from backend.psur.context import PSURContext
        from backend.psur.extraction import extract_from_file

        session = db.query(PSURSession).filter(PSURSession.id == session_id).first()
        if not session:
            raise HTTPException(status_code=404, detail="Session not found")

        sections = db.query(SectionDocument)\
            .filter(SectionDocument.session_id == session_id)\
            .order_by(SectionDocument.section_id)\
            .all()

        # Load template
        template_id = getattr(session, "template_id", None) or "eu_uk_mdr"
        template = load_template(template_id)

        # Try to load PSURContext from saved snapshot first (exact same data as workflow)
        ctx = None
        snapshot_json = getattr(session, "context_snapshot", None) or ""
        if snapshot_json:
            try:
                from dataclasses import fields as dc_fields
                snapshot = json.loads(snapshot_json)
                ctx = PSURContext()
                for f in dc_fields(ctx):
                    if f.name in snapshot:
                        val = snapshot[f.name]
                        # Restore date objects
                        if f.name in ("period_start", "period_end") and isinstance(val, str):
                            try:
                                val = datetime.fromisoformat(val)
                            except Exception:
                                pass
                        # Restore int keys for year dicts
                        if f.name in ("total_units_by_year", "total_complaints_by_year",
                                      "complaint_rate_by_year", "annual_units_golden") and isinstance(val, dict):
                            val = {int(k): v for k, v in val.items()}
                        try:
                            setattr(ctx, f.name, val)
                        except Exception:
                            pass
                print(f"[download] Loaded context from snapshot ({len(snapshot_json)} bytes)")
            except Exception as e:
                print(f"[download] Failed to load context snapshot: {e}. Re-extracting...")
                ctx = None

        # Fallback: reconstruct from files if no snapshot
        if ctx is None:
            print("[download] No context snapshot. Reconstructing from files...")
            ctx = PSURContext(
                device_name=session.device_name or "Unknown Device",
                udi_di=session.udi_di or "",
                period_start=session.period_start,
                period_end=session.period_end,
                template_id=template_id,
            )
            _master = getattr(session, "master_context", None) or {}
            if not isinstance(_master, dict):
                _master = {}
            if _master:
                ctx.manufacturer = str(_master.get("manufacturer", "") or "")
                ctx.manufacturer_address = str(_master.get("manufacturer_address", "") or "")
                ctx.manufacturer_srn = str(_master.get("manufacturer_srn", "") or "")
                ctx.authorized_rep = str(_master.get("authorized_rep", "") or "")
                ctx.notified_body = str(_master.get("notified_body", "") or "")
                ctx.notified_body_number = str(_master.get("notified_body_number", "") or "")
                ctx.intended_use = str(_master.get("intended_use", "") or "")
                ctx.device_type = str(_master.get("device_type", "") or "")

            data_files = db.query(DataFile).filter(DataFile.session_id == session_id).all()
            for df_obj in data_files:
                _file_type = getattr(df_obj, "file_type", "") or ""
                _filename = getattr(df_obj, "filename", "") or ""
                _file_data = getattr(df_obj, "file_data", b"") or b""
                try:
                    extract_from_file(_file_data, _filename, _file_type, ctx)
                except Exception as ext_err:
                    print(f"[download] Extraction error for {_filename}: {ext_err}")

            if _master:
                ed = int(_master.get("exposure_denominator_value", 0) or 0)
                if ed > 0:
                    ctx.total_units_sold = ed
                au = _master.get("annual_units_canonical") or {}
                if au:
                    ctx.total_units_by_year = {int(k): int(v) for k, v in au.items()}
                cc = int(_master.get("complaints_closed_canonical", 0) or 0)
                if cc > 0:
                    ctx.complaints_closed_count = cc

            ctx.calculate_metrics()
            ctx.psur_sequence_number = session_id

        # ---- Build DOCX ----
        doc = DocxDocument()

        # Normal style
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        # === COVER PAGE ===
        doc.add_heading('PERIODIC SAFETY UPDATE REPORT (PSUR)', 0)
        p = doc.add_paragraph()
        p.add_run(f"Device: {session.device_name}").bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p2 = doc.add_paragraph()
        p2.add_run(f"UDI-DI: {session.udi_di or 'Pending'}")
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p3 = doc.add_paragraph()
        period_str = f"{session.period_start.strftime('%d %B %Y') if session.period_start else 'N/A'} to {session.period_end.strftime('%d %B %Y') if session.period_end else 'N/A'}"
        p3.add_run(f"Reporting Period: {period_str}")
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p4 = doc.add_paragraph()
        p4.add_run(f"Regulatory Framework: {template.name}")
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p5 = doc.add_paragraph()
        p5.add_run(f"Generated: {datetime.utcnow().strftime('%d %B %Y')} UTC")
        p5.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Cover page tables
        build_cover_manufacturer_table(doc, ctx)
        build_cover_regulatory_table(doc, ctx, template)
        build_cover_document_table(doc, ctx)

        doc.add_page_break()

        # === TABLE OF CONTENTS placeholder ===
        doc.add_heading("Table of Contents", level=1)
        for sec in sections:
            spec = template.section_specs.get(sec.section_id)
            title = spec.title if spec else sec.section_name
            doc.add_paragraph(
                f"Section {sec.section_id}: {title}",
                style="List Number"
            )
        doc.add_page_break()

        # === SECTION CONTENT ===
        charts = db.query(ChartAsset).filter(ChartAsset.session_id == session_id).all()

        for section in sections:
            spec = template.section_specs.get(section.section_id)
            title = spec.title if spec else section.section_name
            doc.add_heading(f"Section {section.section_id}: {title}", level=1)

            # Insert data tables BEFORE narrative
            print(f"[download] Building tables for section {section.section_id}")
            build_tables_for_section(doc, section.section_id, ctx, template)

            # Insert charts inline AFTER data tables but BEFORE narrative
            section_charts = [c for c in charts if c.section_id == section.section_id]
            for chart in section_charts:
                try:
                    p_title = doc.add_paragraph()
                    run = p_title.add_run(chart.title)
                    run.bold = True
                    run.font.size = Pt(10)
                    run.font.name = 'Calibri'
                    chart_stream = io.BytesIO(chart.png_data)
                    doc.add_picture(chart_stream, width=Inches(5.5))
                    last_para = doc.paragraphs[-1]
                    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as chart_err:
                    doc.add_paragraph(f"[Chart: {chart.title} - could not embed: {chart_err}]")

            # Narrative content with markdown-to-DOCX conversion
            content = section.content or "[No content]"

            # Remove leaked thinking blocks
            if "</thinking>" in content:
                content = content.split("</thinking>")[-1].strip()

            lines = content.split('\n')
            i = 0
            while i < len(lines):
                line = lines[i].strip()

                # Skip empty lines
                if not line:
                    i += 1
                    continue

                # Detect markdown table blocks (consecutive lines with |)
                if '|' in line and line.startswith('|'):
                    table_lines = []
                    while i < len(lines) and '|' in lines[i].strip():
                        table_lines.append(lines[i])
                        i += 1
                    parsed = parse_markdown_table(table_lines)
                    if parsed:
                        insert_markdown_table(doc, parsed)
                    else:
                        for tl in table_lines:
                            doc.add_paragraph(tl.strip())
                    continue

                # Headings
                if line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('# '):
                    doc.add_heading(line[2:], level=2)
                elif line.startswith('- ') or line.startswith('* '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                else:
                    # Regular paragraph with bold/italic support
                    p = doc.add_paragraph()
                    # Parse inline bold (**text**) and italic (*text*)
                    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', line)
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            run = p.add_run(part[2:-2])
                            run.bold = True
                        elif part.startswith('*') and part.endswith('*'):
                            run = p.add_run(part[1:-1])
                            run.italic = True
                        else:
                            p.add_run(part)

                i += 1

        # Save to temp file
        temp_dir = tempfile.gettempdir()
        filename = f"PSUR_{session.device_name.replace(' ', '_')}_{session_id}.docx"
        filepath = os.path.join(temp_dir, filename)
        doc.save(filepath)

        return FileResponse(
            path=filepath,
            filename=filename,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except HTTPException:
        raise
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/sessions/{session_id}/workflow")
async def get_workflow(session_id: int, db: Session = Depends(get_db)):
    """Get workflow state -- returns safe defaults when workflow has not started."""
    try:
        workflow = db.query(WorkflowState)\
            .filter(WorkflowState.session_id == session_id)\
            .first()
        
        if not workflow:
            return {
                "current_section": None,
                "sections_completed": 0,
                "total_sections": 13,
                "status": "not_started",
                "summary": None,
                "paused": False,
                "current_agent": None
            }
        
        return {
            "current_section": getattr(workflow, "current_section", None),
            "sections_completed": getattr(workflow, "sections_completed", 0),
            "total_sections": getattr(workflow, "total_sections", 13),
            "status": getattr(workflow, "status", "not_started"),
            "summary": getattr(workflow, "summary", None),
            "paused": getattr(workflow, "paused", False),
            "current_agent": getattr(workflow, "current_agent", None)
        }
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

# ============================================================================
# SECTION PREVIEW ENDPOINT
# ============================================================================

@app.get("/api/sessions/{session_id}/sections/{section_id}/preview")
async def preview_section(session_id: int, section_id: str, db: Session = Depends(get_db)):
    """Get an HTML preview of a single section with its tables and charts."""
    try:
        section = db.query(SectionDocument).filter(
            SectionDocument.session_id == session_id,
            SectionDocument.section_id == section_id,
        ).first()
        if not section:
            raise HTTPException(status_code=404, detail="Section not found")

        import base64
        from backend.psur.templates import load_template

        template_id = "eu_uk_mdr"
        sess = db.query(PSURSession).filter(PSURSession.id == session_id).first()
        if sess:
            template_id = getattr(sess, "template_id", None) or "eu_uk_mdr"
        template = load_template(template_id)
        spec = template.section_specs.get(section_id)
        title = spec.title if spec else (section.section_name or f"Section {section_id}")

        # Build HTML preview
        html_parts = [
            f"<h2>Section {section_id}: {title}</h2>",
            f"<p style='color:#888;font-size:0.85em;'>Author: {section.author_agent} | "
            f"Status: {section.status} | "
            f"Words: {len(section.content.split()) if section.content else 0}</p>",
        ]

        # Embed charts for this section
        charts = db.query(ChartAsset).filter(
            ChartAsset.session_id == session_id,
            ChartAsset.section_id == section_id,
        ).all()
        for chart in charts:
            b64 = base64.b64encode(chart.png_data).decode("utf-8")
            html_parts.append(
                f"<div style='text-align:center;margin:16px 0;'>"
                f"<p style='font-weight:bold;'>{chart.title}</p>"
                f"<img src='data:image/png;base64,{b64}' style='max-width:100%;'/>"
                f"</div>"
            )

        # Convert markdown content to HTML
        content = section.content or "[No content generated yet]"
        # Remove leaked thinking blocks
        if "</thinking>" in content:
            content = content.split("</thinking>")[-1].strip()

        import re
        for line in content.split("\n"):
            line = line.strip()
            if not line:
                continue
            if line.startswith("### "):
                html_parts.append(f"<h4>{line[4:]}</h4>")
            elif line.startswith("## "):
                html_parts.append(f"<h3>{line[3:]}</h3>")
            elif line.startswith("# "):
                html_parts.append(f"<h3>{line[2:]}</h3>")
            else:
                # Inline bold/italic
                line = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', line)
                line = re.sub(r'\*(.+?)\*', r'<em>\1</em>', line)
                html_parts.append(f"<p>{line}</p>")

        return {
            "section_id": section_id,
            "title": title,
            "status": section.status,
            "word_count": len(section.content.split()) if section.content else 0,
            "charts_count": len(charts),
            "html": "\n".join(html_parts),
        }
    except HTTPException:
        raise
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


# ============================================================================
# CHART ENDPOINTS
# ============================================================================

@app.get("/api/sessions/{session_id}/charts")
async def get_charts(session_id: int, db: Session = Depends(get_db)):
    """Get all generated charts for a session."""
    try:
        charts = db.query(ChartAsset).filter(
            ChartAsset.session_id == session_id
        ).order_by(ChartAsset.chart_id).all()
        return [
            {
                "id": c.id,
                "chart_id": c.chart_id,
                "title": c.title,
                "category": c.category,
                "section_id": c.section_id,
                "created_at": c.created_at.isoformat() if c.created_at else None,
            }
            for c in charts
        ]
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/sessions/{session_id}/charts/{chart_id}")
async def get_chart_image(session_id: int, chart_id: str, db: Session = Depends(get_db)):
    """Get a specific chart as base64 PNG."""
    try:
        chart = db.query(ChartAsset).filter(
            ChartAsset.session_id == session_id,
            ChartAsset.chart_id == chart_id,
        ).first()
        if not chart:
            raise HTTPException(status_code=404, detail="Chart not found")
        import base64
        return {
            "chart_id": chart.chart_id,
            "title": chart.title,
            "category": chart.category,
            "section_id": chart.section_id,
            "base64_png": base64.b64encode(chart.png_data).decode("utf-8"),
        }
    except HTTPException:
        raise
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/sessions/{session_id}/charts/{chart_id}/png")
async def get_chart_png(session_id: int, chart_id: str, db: Session = Depends(get_db)):
    """Get a chart as raw PNG binary (for <img> tags)."""
    try:
        from fastapi.responses import Response
        chart = db.query(ChartAsset).filter(
            ChartAsset.session_id == session_id,
            ChartAsset.chart_id == chart_id,
        ).first()
        if not chart:
            raise HTTPException(status_code=404, detail="Chart not found")
        return Response(content=chart.png_data, media_type="image/png")
    except HTTPException:
        raise
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


# ============================================================================
# WEBSOCKET ENDPOINT
# ============================================================================

@app.websocket("/ws/{session_id}")
async def websocket_endpoint(websocket: WebSocket, session_id: int):
    """WebSocket for real-time updates"""
    await manager.connect(websocket)
    
    try:
        while True:
            # Keep connection alive
            data = await websocket.receive_text()
            
            # Echo back for heartbeat
            await websocket.send_json({
                "type": "heartbeat",
                "timestamp": datetime.utcnow().isoformat()
            })
    except WebSocketDisconnect:
        manager.disconnect(websocket)

# ============================================================================
# BACKGROUND TASKS
# ============================================================================

async def run_orchestrator(session_id: int):
    """Run SOTA orchestrator agent in background"""
    try:
        print(f"\nStarting SOTA orchestrator for session {session_id}...")
        orchestrator = SOTAOrchestrator(session_id)
        
        # Register orchestrator for pause/resume/ask functionality
        active_orchestrators[session_id] = orchestrator

        # Broadcast start
        await manager.broadcast({
            "type": "orchestrator_started",
            "session_id": session_id,
            "timestamp": datetime.utcnow().isoformat()
        })

        print(f"Executing SOTA workflow for session {session_id}...")
        # Execute workflow
        result = await orchestrator.execute_workflow()

        print(f"Workflow complete for session {session_id}: {result}")
        # Broadcast completion
        await manager.broadcast({
            "type": "orchestrator_complete",
            "session_id": session_id,
            "result": result,
            "timestamp": datetime.utcnow().isoformat()
        })

    except Exception as e:
        print(f"Orchestrator error for session {session_id}: {e}")
        traceback.print_exc()
        await manager.broadcast({
            "type": "error",
            "session_id": session_id,
            "error": str(e),
            "timestamp": datetime.utcnow().isoformat()
        })
    finally:
        # Clean up orchestrator reference
        if session_id in active_orchestrators:
            del active_orchestrators[session_id]

# Expose manager for use by agents
def broadcast_message(message: dict):
    """Helper function for agents to broadcast messages"""
    asyncio.create_task(manager.broadcast(message))
