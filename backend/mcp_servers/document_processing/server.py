"""
MCP Document Processing Server
Handles document parsing, template processing, and PSUR assembly
"""

from typing import Dict, List, Any, Optional
from mcp.server import Server
import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
from datetime import datetime


class DocumentProcessingServer:
    """MCP Server for document processing operations"""
    
    def __init__(self):
        self.server = Server("document-processing")
        self.register_tools()
    
    def register_tools(self):
        """Register all document processing tools"""
        
        @self.server.tool()
        async def parse_template(
            template_path: str
        ) -> Dict[str, Any]:
            """
            Parse a PSUR template (FormQAR-054) and extract structure
            
            Args:
                template_path: Path to .docx template file
            
            Returns:
                Dict with template structure, sections, and placeholders
            """
            try:
                doc = docx.Document(template_path)
                
                sections = []
                placeholders = []
                
                for para in doc.paragraphs:
                    text = para.text.strip()
                    
                    # Identify section headers (usually bold or specific style)
                    if para.style.name.startswith('Heading'):
                        sections.append({
                            'text': text,
                            'level': int(para.style.name[-1]) if para.style.name[-1].isdigit() else 1,
                            'style': para.style.name
                        })
                    
                    # Identify placeholders (text in {{brackets}})
                    if '{{' in text and '}}' in text:
                        start = text.find('{{')
                        end = text.find('}}')
                        placeholder = text[start+2:end].strip()
                        placeholders.append(placeholder)
                
                # Count tables
                table_count = len(doc.tables)
                
                return {
                    "status": "success",
                    "template_path": template_path,
                    "sections_found": len(sections),
                    "sections": sections[:20],  # Limit to first 20
                    "placeholders": list(set(placeholders)),
                    "table_count": table_count,
                    "total_paragraphs": len(doc.paragraphs)
                }
                
            except Exception as e:
                return {
                    "error": f"Failed to parse template: {str(e)}",
                    "status": "error"
                }
        
        @self.server.tool()
        async def create_psur_document(
            session_id: int,
            device_name: str,
            output_path: str
        ) -> Dict[str, Any]:
            """
            Create a new PSUR document from template
            
            Args:
                session_id: PSUR session ID
                device_name: Device name for header
                output_path: Where to save the document
            
            Returns:
                Dict with document path and creation status
            """
            try:
                # Create new document
                doc = docx.Document()
                
                # Set up styles
                style = doc.styles['Normal']
                font = style.font
                font.name = 'Calibri'
                font.size = Pt(11)
                
                # Add title
                title = doc.add_paragraph()
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_run = title.add_run(f'Periodic Safety Update Report (PSUR)\n{device_name}')
                title_run.font.size = Pt(16)
                title_run.font.bold = True
                
                # Add metadata
                doc.add_paragraph()  # Spacing
                metadata = doc.add_paragraph()
                metadata.add_run(f'Generated: {datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")}\n')
                metadata.add_run(f'Session ID: {session_id}\n')
                metadata.add_run(f'Multi-Agent PSUR System v1.0')
                
                # Add page break
                doc.add_page_break()
                
                # Save document
                doc.save(output_path)
                
                return {
                    "status": "success",
                    "output_path": output_path,
                    "message": "PSUR document created successfully"
                }
                
            except Exception as e:
                return {
                    "error": f"Failed to create document: {str(e)}",
                    "status": "error"
                }
        
        @self.server.tool()
        async def add_section_to_document(
            document_path: str,
            section_id: str,
            section_title: str,
            content: str,
            author: str
        ) -> Dict[str, Any]:
            """
            Add a section to an existing PSUR document
            
            Args:
                document_path: Path to existing .docx file
                section_id: Section identifier (A, B, C, etc.)
                section_title: Section title
                content: Section content (markdown-formatted)
                author: Agent who authored this section
            
            Returns:
                Dict with addition status
            """
            try:
                # Open existing document
                doc = docx.Document(document_path)
                
                # Add section header
                heading = doc.add_heading(f'Section {section_id}: {section_title}', level=1)
                
                # Add author metadata (small, gray text)
                author_para = doc.add_paragraph()
                author_run = author_para.add_run(f'Generated by: {author}')
                author_run.font.size = Pt(9)
                author_run.font.color.rgb = RGBColor(128, 128, 128)
                author_para.add_run(f' • {datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")}')
                
                # Parse and add content
                lines = content.split('\n')
                for line in lines:
                    line = line.strip()
                    
                    if not line:
                        continue
                    
                    # Check if it's a heading
                    if line.startswith('###'):
                        doc.add_heading(line.replace('###', '').strip(), level=3)
                    elif line.startswith('##'):
                        doc.add_heading(line.replace('##', '').strip(), level=2)
                    elif line.startswith('**') and line.endswith('**'):
                        # Bold paragraph
                        para = doc.add_paragraph()
                        run = para.add_run(line.strip('*'))
                        run.bold = True
                    else:
                        # Normal paragraph
                        doc.add_paragraph(line)
                
                # Add page break after section
                doc.add_page_break()
                
                # Save document
                doc.save(document_path)
                
                return {
                    "status": "success",
                    "section_added": section_id,
                    "author": author,
                    "message": f"Section {section_id} added successfully"
                }
                
            except Exception as e:
                return {
                    "error": f"Failed to add section: {str(e)}",
                    "status": "error"
                }
        
        @self.server.tool()
        async def extract_data_from_file(
            file_path: str,
            file_type: str
        ) -> Dict[str, Any]:
            """
            Extract data from uploaded files (Excel, CSV, Word)
            
            Args:
                file_path: Path to file
                file_type: Type of file (excel, csv, word, pdf)
            
            Returns:
                Dict with extracted data or summary
            """
            try:
                if file_type in ['excel', 'xlsx', 'xls']:
                    # Read Excel file
                    df = pd.read_excel(file_path)
                    
                    return {
                        "status": "success",
                        "file_type": "excel",
                        "rows": len(df),
                        "columns": list(df.columns),
                        "preview": df.head(5).to_dict('records'),
                        "summary": {
                            "total_rows": len(df),
                            "total_columns": len(df.columns),
                            "missing_values": df.isnull().sum().to_dict()
                        }
                    }
                
                elif file_type in ['csv']:
                    # Read CSV file
                    df = pd.read_csv(file_path)
                    
                    return {
                        "status": "success",
                        "file_type": "csv",
                        "rows": len(df),
                        "columns": list(df.columns),
                        "preview": df.head(5).to_dict('records')
                    }
                
                elif file_type in ['docx', 'word']:
                    # Read Word file
                    doc = docx.Document(file_path)
                    
                    content = []
                    for para in doc.paragraphs:
                        if para.text.strip():
                            content.append(para.text.strip())
                    
                    return {
                        "status": "success",
                        "file_type": "word",
                        "paragraphs": len(content),
                        "content_preview": content[:10],
                        "table_count": len(doc.tables)
                    }
                
                else:
                    return {
                        "error": f"Unsupported file type: {file_type}",
                        "status": "error"
                    }
                
            except Exception as e:
                return {
                    "error": f"Failed to extract data: {str(e)}",
                    "status": "error"
                }
    
    async def start(self, host: str = "localhost", port: int = 8002):
        """Start the MCP server"""
        await self.server.run(host=host, port=port)


# Server instance
document_processing_server = DocumentProcessingServer()


if __name__ == "__main__":
    import asyncio
    asyncio.run(document_processing_server.start())
